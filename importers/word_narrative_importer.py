import os
import re
from datetime import datetime
from xml.etree.ElementTree import XML

from docx import Document

from pepys_import.core.validators import constants
from pepys_import.file.importer import Importer

WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
TEXT = WORD_NAMESPACE + "t"


class WordNarrativeImporter(Importer):
    def __init__(self):
        super().__init__(
            name="Word Narrative Format Importer",
            validation_level=constants.BASIC_LEVEL,
            short_name="Word Narrative Importer",
            default_privacy="Public",
            datafile_type="Word Narrative",
        )

        self.last_day = None
        self.last_month = None
        self.last_year = None

    def can_load_this_type(self, suffix):
        return suffix.upper() in [".DOCX", ".PDF"]

    def can_load_this_filename(self, filename):
        return True

    def can_load_this_header(self, header):
        return True

    def can_load_this_file(self, file_contents):
        return True

    def _load_this_file(self, data_store, path, file_object, datafile, change_id):
        _, ext = os.path.splitext(path)
        if ext.upper() == ".DOCX":
            header, entries, error = self.load_docx_file(path)
        elif ext.upper() == ".PDF":
            header, entries, error = self.load_pdf_file(path)
        else:
            self.errors.append({self.error_type: f"Unsupported file extension {ext}."})
            return

        if error:
            # Stop parsing if there was an error during loading that we can't recover from
            return

        self.parse_file(header, entries, data_store, change_id)

    def parse_file(self, header, entries, data_store, change_id):
        platform_from_header = header.get("platform", None)
        platform = self.get_cached_platform(
            data_store, platform_name=platform_from_header, change_id=change_id
        )
        print(platform)

        # Loop through each entry in the file
        for entry in entries:
            stripped_entry = entry.strip()
            print(f"Entry {stripped_entry}")
            if stripped_entry == "":
                # Skip blank entries
                continue

            parts = stripped_entry.split(",")

            correct_length = len(parts) > 5
            has_length_and_four_fig_datetime = correct_length and re.fullmatch(r"\d{4}", parts[0])
            has_length_and_six_fig_datetime = correct_length and re.fullmatch(r"\d{6}", parts[0])

            is_comma_sep_with_datetime = (
                has_length_and_four_fig_datetime or has_length_and_six_fig_datetime
            )

            if is_comma_sep_with_datetime:
                self.process_comma_sep_entry(header, parts, has_length_and_four_fig_datetime)
            else:
                # The entry isn't comma separated with a datetime at the start
                # These entries mostly occur in PDFs not DOCXs - but we check for them
                # everywhere.
                # Even though it isn't comma separated, it might still have a date at the
                # beginning and look like this:
                # 120500 Message 1 (NB: the message could still include FCS entries etc)
                # Or it could be a date block marker like this:
                # 12 Dec 95
                # Or it could be a bit of text that just needs adding on to the previous entry
                # So, check for these one at a time
                #
                # Here we check if it starts with 4 or 6 digits, followed by whitespace
                if re.match(r"\d{4}\w", stripped_entry) or re.match(r"\d{6}\w", stripped_entry):
                    # If so, we process the entry
                    self.process_non_comma_entry(header, stripped_entry)
                else:
                    # Try parsing the line as a date in the formats
                    # dd MMM yy
                    # dd MMM yyyy
                    # For example, "12 DEC 1995"
                    formats = ["%d %b %y", "%d %b %Y"]
                    timestamp = None
                    for date_format in formats:
                        try:
                            timestamp = datetime.strptime(stripped_entry, date_format)
                        except ValueError:
                            continue

                    if timestamp is not None:
                        # We've got a valid timestamp
                        # So store the details ready for use with any lines that follow it
                        self.last_day = timestamp.day
                        self.last_month = timestamp.month
                        self.last_year = timestamp.year
                        continue

                    # If we've got here, then we just have some text that needs appending to the previous entry
                    # TODO: Append entry

    def process_non_comma_entry(self, header, stripped_entry):
        print(f"Found non comma entry: {stripped_entry}")
        split_by_whitespace = stripped_entry.split()
        timestamp_str = split_by_whitespace[0].trim()

        try:
            timestamp = self.parse_singlepart_datetime(timestamp_str)
        except Exception as e:
            self.errors.append(
                {self.error_type: f"Error parsing timestamp {timestamp_str}, error was {str(e)}"}
            )
            return

        print(timestamp)

    def parse_singlepart_datetime(self, timestamp_str):
        if self.last_day is None or self.last_month is None or self.last_year is None:
            raise ValueError("No previous day/month/year block")

        if len(timestamp_str) == 6:
            day = int(timestamp_str[0:2])
            hour = int(timestamp_str[2:4])
            mins = int(timestamp_str[4:6])

            if day < self.last_day:
                # Day has gone down, so month must go up

                # However, if month is 12 then it must go to 1 and year must go up
                if self.last_month == 12:
                    month = 1
                    year = self.last_year + 1
                else:
                    month = self.last_month + 1
                    year = self.last_year
            else:
                month = self.last_month
                year = self.last_year

            timestamp = datetime(year, month, day, hour, mins)
            return timestamp
        elif len(timestamp_str) == 4:
            hour = int(timestamp_str[0:2])
            mins = int(timestamp_str[2:4])

            timestamp = datetime(self.last_year, self.last_month, self.last_day, hour, mins)
            return timestamp
        else:
            raise ValueError("Timestamp must be 4 digits (HHMM) or 6 digits (DDHHMM)")

    def process_comma_sep_entry(self, header, parts, has_length_and_four_fig_datetime):
        # Parse datetime
        timestamp, error = self.parse_multipart_datetime(
            parts, four_fig=has_length_and_four_fig_datetime
        )
        if error:
            return

        # Process rest of entry
        entry_platform_name = parts[4].strip()

        if entry_platform_name.upper() != header["platform"].upper():
            header_platform_name = header["platform"]
            self.errors.append(
                {
                    self.error_type: f"Platform name in entry ('{entry_platform_name}') doesn't match platform name in header ('{header_platform_name}')"
                }
            )
            return

        message_type = parts[5].strip()

        if message_type.upper() == "FCS":
            # It's a Fire Control Solution message
            self.process_fcs_message(timestamp, entry_platform_name, parts[6:])
        else:
            # It's another type of message
            if len(message_type) > 20:
                # Sometimes there isn't the end comma on the message type field
                # which means it gets merged with the text field
                # If this field is very long then this is probably what happened
                # So we find the first location of a tab, and split on that
                index = message_type.find("\t")
                if index != -1:
                    text = message_type[index:].strip()
                    message_type = message_type[:index].strip()
                else:
                    fulltext = ",".join(parts)
                    self.errors.append(
                        {
                            self.error_type: f"Can't separate message type and text, are fields mangled or a comma missing? {fulltext}"
                        }
                    )
                return
            else:
                text = ",".join(parts[6:]).strip()

            print(f"Timestamp: {timestamp}")
            print(f"message_type: {message_type}")
            print(f"text: {text}")

            # TODO: Work out here if we've got a state entry in the comment
            # and if so then parse it and store it

            # Store message data here
            self.store_comment(timestamp, entry_platform_name, message_type, text)

    def process_fcs_message(self, timestamp, platform_name, fcs_parts):
        pass

    def store_comment(self, timestamp, entry_platform_name, message_type, text):
        pass

    def parse_multipart_datetime(self, parts, four_fig):
        day_visible = None

        # Get the parts separated by commas, as they're always there
        day_hidden = int(parts[1])
        month = int(parts[2])
        year = int(parts[3])

        if four_fig:
            # It's a four figure time with just HHMM
            hour = int(parts[0][0:2])
            mins = int(parts[0][2:4])
        else:
            # It's a six figure time with DDHHMM
            day_visible = int(parts[0][0:2])  # day in the visible text
            hour = int(parts[0][2:4])
            mins = int(parts[0][4:6])

            # Deal with entries that might need to be pulled back from the next day
            # If something that happened at 2345 only gets entered at 0010 then
            # the hidden text will have the next day in it, when it should be
            # the previous day
            if hour == 23:
                if day_hidden == day_visible + 1:
                    day_hidden = day_visible

            if day_hidden != day_visible:
                full_text = ",".join(parts)
                self.errors.append(
                    {
                        self.error_type: f"Day in text doesn't match day in hidden text - possible copy/paste error: '{full_text}'."
                    }
                )
                return None, True

        day = day_visible or day_hidden

        day_decreased = (self.last_day is not None) and (day < self.last_day)
        month_increased = (self.last_month is not None) and (month > self.last_month)
        year_increased = (self.last_month is not None) and (year > self.last_year)

        # Deal with entries where the day has decreased (ie. gone to the beginning of the next month)
        # but the month and/or year hasn't increased
        # This suggests that there has been a copy-paste error, mangling the data
        if day_decreased and ((not month_increased) or (not year_increased)):
            self.errors.append(
                {self.error_type: f"Day decreased but month/year didn't increase: {parts[0]}."}
            )
            return None, True
        else:
            # Everything makes sense, so we can update the last_X variables
            self.last_day = day_visible
            self.last_month = month
            self.last_year = year

        if year < 100:
            # If a two digit year
            if year > 80:
                # If it is from 80s onwards then it's 1900s
                year = 1900 + year
            else:
                year = 2000 + year

        if year < 1900 or year > 2100:
            self.errors.append({self.error_type: f"Year too big or too small: {year}."})
            return None, True

        try:
            timestamp = datetime(year, month, day, hour, mins)
        except ValueError:
            full_text = ",".join(parts)
            self.errors.append({self.error_type: f"Could not parse timestamp {full_text}."})
            return None, True

        return timestamp, False

    def load_docx_file(self, path):
        try:
            doc = Document(path)
        except Exception as e:
            self.errors.append(
                {self.error_type: f'Invalid docx file at {path}\nError from parsing was "{str(e)}"'}
            )
            return None, None, True

        try:
            # Get text from the header
            # Headers are attached to a document section, so we need to extract the section first
            sec = doc.sections[0]
            header_text = ""
            for para in sec.header.paragraphs:
                header_text += "\n" + para.text

            splitted = re.split("[\n\t]+", header_text.strip())
            header = {}
            header["privacy"] = splitted[0].strip()
            header["platform"] = splitted[1].strip()
            header["exercise"] = splitted[4].strip()
            header["fulltext"] = header_text.strip()
        except Exception:
            # Couldn't extract header, so presumably doesn't have a header
            # That's ok - we just create an empty dict
            header = {}

        try:
            # Get each paragraph entry, after accepting any tracked changes
            entries = []
            for p in doc.paragraphs:
                entries.append(self.get_accepted_text(p))
        except Exception as e:
            self.errors.append(
                {self.error_type: f'Cannot extract paragraphs\nError from parsing was "{str(e)}"'}
            )
            return None, None, True

        return header, entries, False

    def get_accepted_text(self, p):
        """Return text of a paragraph after accepting all changes.

        This gets the XML content of the paragraph and checks for deletions or insertions. If there
        aren't any, then it just returns the text. If there are some, then it parses the XML and
        joins the individual text entries."""
        # Taken from https://stackoverflow.com/questions/38247251/how-to-extract-text-inserted-with-track-changes-in-python-docx
        xml = p._p.xml
        if "w:del" in xml or "w:ins" in xml:
            tree = XML(xml)
            runs = (node.text for node in tree.iter(TEXT) if node.text)
            return "".join(runs)
        else:
            return p.text
